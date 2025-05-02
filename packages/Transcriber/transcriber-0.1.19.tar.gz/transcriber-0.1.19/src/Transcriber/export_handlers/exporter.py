import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt

from Transcriber.config import settings
from Transcriber.types.export_type import ExportType
from Transcriber.types.segment_type import SegmentType
from Transcriber.utils import time_utils


class Writer:
    def write_all(self, file_name: str, segments: list[SegmentType]) -> None:
        should_compact = settings.output.min_words_per_segment > 0
        segments_to_write = (
            self.compact_segments(segments, settings.output.min_words_per_segment) if should_compact else segments
        )
        for output_format in settings.output.output_formats:
            format_output_dir = os.path.join(settings.output.output_dir, output_format)
            # Save original if needed
            if settings.output.save_files_before_compact and should_compact:
                self.write(
                    ExportType(output_format),
                    format_output_dir,
                    f"{file_name}-original.{output_format}",
                    segments,
                )

            # Save compacted or regular version
            self.write(
                ExportType(output_format),
                format_output_dir,
                f"{file_name}.{output_format}",
                segments_to_write,
            )

    def write(
        self,
        export_format: ExportType,
        output_dir: str,
        file_path: str,
        segments: list[SegmentType],
    ) -> None:
        file_path = os.path.join(output_dir, file_path)
        if export_format == ExportType.TXT:
            self.write_txt(file_path, segments)
        elif export_format == ExportType.SRT:
            self.write_srt(file_path, segments)
        elif export_format == ExportType.VTT:
            self.write_vtt(file_path, segments)
        elif export_format == ExportType.DOCX:
            self.write_docx(file_path, segments)

    def write_txt(self, file_path: str, segments: list[SegmentType]) -> None:
        self._write_to_file(file_path, self.generate_txt(segments))

    def write_srt(self, file_path: str, segments: list[SegmentType]) -> None:
        self._write_to_file(file_path, self.generate_srt(segments))

    def write_vtt(self, file_path: str, segments: list[SegmentType]) -> None:
        self._write_to_file(file_path, self.generate_vtt(segments))

    def write_docx(self, file_path: str, segments: list[SegmentType]) -> None:
        doc = Document()
        file_name = os.path.basename(file_path)
        title = os.path.splitext(file_name)[0]
        header = doc.add_heading(title, level=1)

        if self.is_rtl():
            # Set the header direction to RTL
            self.set_element_rtl(header)

        # Set the header alignment to center
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Set the font size and name for the header
        header_font = header.runs[0].font
        header_font.size = Pt(settings.output.title_font_size)
        header_font.name = settings.output.title_font_name

        for segment in segments:
            paragraph_text = segment["text"].strip()

            if self.is_rtl():
                # Replace characters for RTL
                paragraph_text = self.prepare_text_for_rtl(paragraph_text)

            paragraph = doc.add_paragraph(paragraph_text)

            if self.is_rtl():
                # Set the paragraph direction to RTL using XML
                self.set_element_rtl(paragraph)
                # In RTL mode, LEFT alignment is actually right-aligned
                # Set the alignment to LEFT, which will be interpreted as right-aligned in RTL mode
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # change font size
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(settings.output.body_font_size)
                font.name = settings.output.body_font_name

        doc.save(file_path)

    def generate_txt(self, segments: list[SegmentType]) -> str:
        return "\n".join(list(map(lambda segment: segment["text"].strip(), segments))) + "\n"

    def generate_srt(self, segments: list[SegmentType]) -> str:
        return "".join(
            f"{i}\n"
            f"{time_utils.format_timestamp(segment['start'], include_hours=True, decimal_marker=',')} --> "
            f"{time_utils.format_timestamp(segment['end'], include_hours=True, decimal_marker=',')}\n"
            f"{segment['text'].strip()}\n\n"
            for i, segment in enumerate(segments, start=1)
        )

    def generate_vtt(self, segments: list[SegmentType]) -> str:
        return "WEBVTT\n\n" + "".join(
            f"{time_utils.format_timestamp(segment['start'])} --> {time_utils.format_timestamp(segment['end'])}\n"
            f"{segment['text'].strip()}\n\n"
            for segment in segments
        )

    def compact_segments(self, segments: list[SegmentType], min_words_per_segment: int) -> list[SegmentType]:
        if min_words_per_segment == 0:
            return segments

        compacted_segments = []
        tmp_segment = None

        for segment in segments:
            if tmp_segment:
                tmp_segment["text"] += f" {segment['text'].strip()}"
                tmp_segment["end"] = segment["end"]

                if len(tmp_segment["text"].split()) >= min_words_per_segment:
                    compacted_segments.append(tmp_segment)
                    tmp_segment = None
            elif len(segment["text"].split()) < min_words_per_segment:
                tmp_segment = segment.copy()
            elif len(segment["text"].split()) >= min_words_per_segment:
                compacted_segments.append(segment.copy())

        if tmp_segment:
            compacted_segments.append(tmp_segment)

        return compacted_segments

    def is_output_exist(self, file_name: str):
        if settings.output.save_files_before_compact and not all(
            Path(
                settings.output.output_dir,
                output_format,
                f"{file_name}-original.{output_format}",
            ).is_file()
            for output_format in settings.output.output_formats
        ):
            return False

        if (not settings.output.save_files_before_compact or settings.output.min_words_per_segment != 0) and not all(
            Path(
                settings.output.output_dir,
                output_format,
                f"{file_name}.{output_format}",
            ).is_file()
            for output_format in settings.output.output_formats
        ):
            return False

        return True

    def _write_to_file(self, file_path: str, content: str) -> None:
        with open(file_path, "w", encoding="utf-8") as file:
            file.write(content)

    def is_rtl(self) -> bool:
        """Checks if the current output setting is right-to-left."""
        if settings.whisper.language == "ar":
            return True
        return False

    def set_element_rtl(self, element):
        """Sets the element to right-to-left using XML."""
        # Get element properties container that holds formatting attributes
        element_properties = element._element.get_or_add_pPr()
        # Create a new XML element for right-to-left direction
        rtl_element = OxmlElement("w:bidi")
        # Append the element to the element properties container
        element_properties.append(rtl_element)

    def prepare_text_for_rtl(self, text: str) -> str:
        """
        Replaces characters in the text from LTR to RTL.
        This is a placeholder function and should be implemented based on specific requirements.
        """
        # Define the replacements for LTR to RTL conversion
        # Convert common punctuation marks
        replacements = {
            "(": "）",
            ")": "（",
            "[": "］",
            "]": "［",
            "{": "｝",
            "}": "｛",
            "<": "＞",
            ">": "＜",
            ",": "،",
            ".": ".",
            "?": "؟",
            "!": "！",
        }
        for ltr, rtl in replacements.items():
            text = text.replace(ltr, rtl)
        return text
