import os
import time
import json
import hashlib
import sys
import asyncio
from pathlib import Path
from typing import List, Optional, Set, Dict, Any, Tuple

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor
from openai import OpenAI
from openai import AsyncOpenAI
from loguru import logger
from tqdm import tqdm

# Constants
DEFAULT_MODEL = "gpt-4-turbo"
MAX_RETRIES = 5
INITIAL_RETRY_DELAY = 1
TEMPERATURE = 0.3
DEFAULT_CACHE_DIR = Path.home() / ".cache/.docx_translator"
DEFAULT_MAX_CONCURRENT = 5
DEFAULT_TARGET_LANGUAGE = os.environ.get("TARGET_LANGUAGE", "Spanish")

# Configure logging with loguru
logger.remove()  # Remove default handler
logger.add(
    sys.stderr,
    format="<green>{time:YYYY-MM-DD HH:mm:ss}</green> | <level>{level: <8}</level> | <level>{message}</level>",
    level="INFO",
    colorize=True,
)
logger.add(
    "docx_translator.log",
    rotation="10 MB",
    retention="1 week",
    level="DEBUG",
    compression="zip",
)

# Create a dictionary to track streaming progress
translation_progress = {}
total_tokens_received = 0
total_cached_tokens = 0  # Track total tokens from cache


def setup_openai_client(
    api_key: Optional[str] = None, base_url: Optional[str] = None
) -> OpenAI:
    """Set up and return OpenAI client

    Args:
        api_key: OpenAI API key (optional if environment variable is set)
        base_url: OpenAI API base URL for custom endpoints (optional)

    Returns:
        OpenAI client instance

    Raises:
        ValueError: If API key is not provided and not found in environment
    """
    key = api_key or os.environ.get("OPENAI_API_KEY")
    if not key:
        raise ValueError(
            "OpenAI API key is required. Set it via --api_key or OPENAI_API_KEY environment variable."
        )

    client_kwargs = {"api_key": key}

    # Add base_url if provided
    if base_url:
        client_kwargs["base_url"] = base_url
        logger.info(f"Using custom OpenAI API base URL: {base_url}")

    return OpenAI(**client_kwargs)


def setup_async_openai_client(
    api_key: Optional[str] = None, base_url: Optional[str] = None
) -> AsyncOpenAI:
    """Set up and return async OpenAI client

    Args:
        api_key: OpenAI API key (optional if environment variable is set)
        base_url: OpenAI API base URL for custom endpoints (optional)

    Returns:
        AsyncOpenAI client instance

    Raises:
        ValueError: If API key is not provided and not found in environment
    """
    key = api_key or os.environ.get("OPENAI_API_KEY")
    if not key:
        raise ValueError(
            "OpenAI API key is required. Set it via --api_key or OPENAI_API_KEY environment variable."
        )

    client_kwargs = {"api_key": key}

    # Add base_url if provided
    if base_url:
        client_kwargs["base_url"] = base_url
        logger.info(f"Using custom OpenAI API base URL: {base_url}")

    return AsyncOpenAI(**client_kwargs)


class TranslationCache:
    """Cache for translations to avoid redundant API calls"""

    def __init__(
        self,
        target_language: str,
        model: str = DEFAULT_MODEL,
        cache_dir: Path = DEFAULT_CACHE_DIR,
    ):
        """Initialize the translation cache

        Args:
            target_language: The target language for translations
            model: The model used for translations
            cache_dir: Directory to store cache files
        """
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True, parents=True)

        # Create a cache file specific to this language and model
        cache_filename = f"cache_{target_language.lower().replace(' ', '_')}_{model.replace('-', '_')}.json"
        self.cache_file = self.cache_dir / cache_filename

        # Load existing cache if available
        self.cache: Dict[str, Dict[str, Any]] = {}
        if self.cache_file.exists():
            try:
                with open(self.cache_file, "r", encoding="utf-8") as f:
                    self.cache = json.load(f)
                logger.info(
                    f"Loaded {len(self.cache)} cached translations for {target_language}"
                )
            except Exception as e:
                logger.warning(f"Failed to load translation cache: {e}")
                # Start with an empty cache if loading fails
                self.cache = {}

    def get_cache_key(self, text: str) -> str:
        """Generate a unique key for the text

        Args:
            text: The text to be cached

        Returns:
            A unique hash for the text
        """
        # Use MD5 hash of the text as the key
        return hashlib.md5(text.encode("utf-8")).hexdigest()

    def get(self, text: str) -> Optional[Dict[str, Any]]:
        """Get a translation from the cache

        Args:
            text: The text to get translation for

        Returns:
            Dictionary with translated text and token usage or None if not found
        """
        key = self.get_cache_key(text)
        return self.cache.get(key)

    def set(
        self, text: str, translation: str, token_usage: Optional[Dict[str, int]] = None
    ) -> None:
        """Store a translation in the cache with token usage information

        Args:
            text: The original text
            translation: The translated text
            token_usage: Dictionary with token usage information
        """
        key = self.get_cache_key(text)

        # Store both the translation and token usage
        self.cache[key] = {
            "translation": translation,
            "token_usage": token_usage or {},
        }

        # Periodically save the cache to disk (every 10 new entries)
        if len(self.cache) % 10 == 0:
            self.save()

    def save(self) -> None:
        """Save the cache to disk"""
        try:
            with open(self.cache_file, "w", encoding="utf-8") as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
            logger.debug(f"Saved translation cache with {len(self.cache)} entries")
        except Exception as e:
            logger.error(f"Failed to save translation cache: {e}")


async def stream_async_translation(
    text: str,
    target_language: str,
    client: AsyncOpenAI,
    cache: Optional[TranslationCache] = None,
    progress_bar: Optional[tqdm] = None,
    task_id: Optional[str] = None,
    cancellation_check=None,
) -> Tuple[str, str]:
    """Translate text asynchronously with streaming using the OpenAI API

    Args:
        text: The text to translate
        target_language: The target language
        client: AsyncOpenAI client instance
        cache: Optional TranslationCache instance
        progress_bar: Optional tqdm progress bar to update
        task_id: Optional task ID for tracking progress
        cancellation_check: Optional function to check if translation should be cancelled

    Returns:
        Tuple containing the translated text and translation status
    """
    global total_tokens_received, total_cached_tokens, translation_progress

    # Check if translation should be cancelled
    if cancellation_check and cancellation_check():
        return "", "cancelled"

    # Initialize task progress if task_id is provided
    if task_id:
        translation_progress[task_id] = {"status": "starting", "text": "", "tokens": 0}

    # Try to get from cache first if cache is provided
    if cache:
        cached_translation = cache.get(text)
        if cached_translation:
            translated_text = cached_translation["translation"]
            # Update token counts if available
            token_usage = cached_translation.get("token_usage", {})
            if token_usage:
                total_cached_tokens += token_usage.get("total_tokens", 0)

            # Update progress
            if task_id:
                translation_progress[task_id] = {
                    "status": "completed",
                    "text": translated_text,
                    "tokens": token_usage.get("total_tokens", 0),
                    "from_cache": True,
                }

            # Update progress bar if provided
            if progress_bar:
                progress_bar.update(1)

            return translated_text, "cached"

    # Set status to in_progress
    if task_id:
        translation_progress[task_id] = {
            "status": "in_progress",
            "text": "",
            "tokens": 0,
        }

    # Prepare the prompt
    prompt = f"Translate the following text into {target_language}. Keep formatting such as bold, italic, or links intact. Translate only the text itself without adding any explanations.\n\nText to translate:\n\n{text}"

    for attempt in range(MAX_RETRIES):
        try:
            # Check for cancellation before making API call
            if cancellation_check and cancellation_check():
                return "", "cancelled"

            # Streaming response through chat completions API
            stream = await client.chat.completions.create(
                model=DEFAULT_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=TEMPERATURE,
                stream=True,
            )

            translated_text = ""
            tokens_count = 0
            async for chunk in stream:
                # Check for cancellation
                if cancellation_check and cancellation_check():
                    # Close the stream if possible
                    try:
                        await stream.aclose()
                    except:
                        pass
                    return "", "cancelled"

                # Get the content from the chunk if available
                content = None
                if (
                    chunk.choices
                    and len(chunk.choices) > 0
                    and chunk.choices[0].delta.content is not None
                ):
                    content = chunk.choices[0].delta.content

                if content:
                    translated_text += content
                    tokens_count += 1
                    if task_id:
                        translation_progress[task_id]["text"] = translated_text
                        translation_progress[task_id]["tokens"] = tokens_count

            # Get token usage if available
            token_usage = None
            try:
                # Estimate token usage based on input and output length
                # Add approximate token counts (rough estimation)
                input_tokens = len(text.split()) * 1.3  # Approximate tokens in input
                output_tokens = (
                    len(translated_text.split()) * 1.3
                )  # Approximate tokens in output
                total_tokens = int(input_tokens + output_tokens)
                token_usage = {"total_tokens": total_tokens}
                total_tokens_received += total_tokens
            except Exception as e:
                logger.warning(f"Could not estimate token usage: {e}")

            # Store in cache
            if cache:
                cache.set(text, translated_text, token_usage)

            # Update progress
            if task_id:
                translation_progress[task_id] = {
                    "status": "completed",
                    "text": translated_text,
                    "tokens": (
                        token_usage.get("total_tokens", 0)
                        if token_usage
                        else tokens_count
                    ),
                    "from_cache": False,
                }

            # Update progress bar if provided
            if progress_bar:
                progress_bar.update(1)

            return translated_text, "completed"

        except Exception as e:
            # Exponential backoff with jitter
            retry_delay = INITIAL_RETRY_DELAY * (2**attempt) + INITIAL_RETRY_DELAY * (
                attempt * 0.1
            )
            logger.exception(e)
            logger.warning(
                f"Translation attempt {attempt + 1}/{MAX_RETRIES} failed: {e}, retrying in {retry_delay:.2f}s"
            )

            # Check for cancellation before sleeping
            if cancellation_check and cancellation_check():
                return "", "cancelled"

            await asyncio.sleep(retry_delay)

            # Check for cancellation again after sleeping
            if cancellation_check and cancellation_check():
                return "", "cancelled"

    # If we're here, all retries failed
    error_message = f"Failed to translate text after {MAX_RETRIES} attempts"
    logger.error(error_message)

    # Update progress with error status
    if task_id:
        translation_progress[task_id] = {
            "status": "error",
            "text": error_message,
            "tokens": 0,
        }

    return "", "error"


def stream_translation(
    text: str,
    target_language: str,
    client: OpenAI,
    cache: Optional[TranslationCache] = None,
    progress_bar: Optional[tqdm] = None,
    task_id: Optional[str] = None,
) -> str:
    """Translate text with streaming using the OpenAI API (synchronous version)

    Args:
        text: The text to translate
        target_language: The target language
        client: OpenAI client instance
        cache: Optional TranslationCache instance
        progress_bar: Optional tqdm progress bar to update
        task_id: Optional task ID for tracking progress

    Returns:
        The translated text
    """
    global total_tokens_received, total_cached_tokens, translation_progress

    # Initialize task progress if task_id is provided
    if task_id:
        translation_progress[task_id] = {"status": "starting", "text": "", "tokens": 0}

    # Try to get from cache first if cache is provided
    if cache:
        cached_translation = cache.get(text)
        if cached_translation:
            translated_text = cached_translation["translation"]
            # Update token counts if available
            token_usage = cached_translation.get("token_usage", {})
            if token_usage:
                total_cached_tokens += token_usage.get("total_tokens", 0)

            # Update progress
            if task_id:
                translation_progress[task_id] = {
                    "status": "completed",
                    "text": translated_text,
                    "tokens": token_usage.get("total_tokens", 0),
                    "from_cache": True,
                }

            # Update progress bar if provided
            if progress_bar:
                progress_bar.update(1)

            return translated_text

    # Set status to in_progress
    if task_id:
        translation_progress[task_id] = {
            "status": "in_progress",
            "text": "",
            "tokens": 0,
        }

    # Prepare the prompt
    prompt = f"Translate the following text into {target_language}. Keep formatting such as bold, italic, or links intact. Translate only the text itself without adding any explanations.\n\nText to translate:\n\n{text}"

    for attempt in range(MAX_RETRIES):
        try:
            # Streaming response through chat completions API
            stream = client.chat.completions.create(
                model=DEFAULT_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=TEMPERATURE,
                stream=True,
            )

            translated_text = ""
            tokens_count = 0
            for chunk in stream:
                # Get the content from the chunk if available
                content = None
                if (
                    chunk.choices
                    and len(chunk.choices) > 0
                    and chunk.choices[0].delta.content is not None
                ):
                    content = chunk.choices[0].delta.content

                if content:
                    translated_text += content
                    tokens_count += 1
                    if task_id:
                        translation_progress[task_id]["text"] = translated_text
                        translation_progress[task_id]["tokens"] = tokens_count

            # Get token usage if available
            token_usage = None
            try:
                # Estimate token usage based on input and output length
                # Add approximate token counts (rough estimation)
                input_tokens = len(text.split()) * 1.3  # Approximate tokens in input
                output_tokens = (
                    len(translated_text.split()) * 1.3
                )  # Approximate tokens in output
                total_tokens = int(input_tokens + output_tokens)
                token_usage = {"total_tokens": total_tokens}
                total_tokens_received += total_tokens
            except Exception as e:
                logger.warning(f"Could not estimate token usage: {e}")

            # Store in cache
            if cache:
                cache.set(text, translated_text, token_usage)

            # Update progress
            if task_id:
                translation_progress[task_id] = {
                    "status": "completed",
                    "text": translated_text,
                    "tokens": (
                        token_usage.get("total_tokens", 0)
                        if token_usage
                        else tokens_count
                    ),
                    "from_cache": False,
                }

            # Update progress bar if provided
            if progress_bar:
                progress_bar.update(1)

            return translated_text

        except Exception as e:
            # Exponential backoff with jitter
            retry_delay = INITIAL_RETRY_DELAY * (2**attempt) + INITIAL_RETRY_DELAY * (
                attempt * 0.1
            )
            logger.exception(e)
            logger.warning(
                f"Translation attempt {attempt + 1}/{MAX_RETRIES} failed: {e}, retrying in {retry_delay:.2f}s"
            )
            time.sleep(retry_delay)

    # If we're here, all retries failed
    error_message = f"Failed to translate text after {MAX_RETRIES} attempts"
    logger.error(error_message)

    # Update progress with error status
    if task_id:
        translation_progress[task_id] = {
            "status": "error",
            "text": error_message,
            "tokens": 0,
        }

    return ""


def add_translation_to_paragraph(paragraph, translated_text: str) -> None:
    """Add translated text to a paragraph with a line break

    Args:
        paragraph: The paragraph object to add translation to
        translated_text: The translated text to add
    """
    # Add a line break and then the separator line
    run = paragraph.add_run()
    run.add_break(WD_BREAK.LINE)  # Add a line break

    # Add separator line
    separator_run = paragraph.add_run("------")
    separator_run.add_break(WD_BREAK.LINE)  # Add another line break after separator

    # Add translation with gray color
    translation_run = paragraph.add_run(translated_text)
    translation_run.font.color.rgb = RGBColor(128, 128, 128)  # Set to gray color


def count_translatable_elements(doc: Document, target_styles_set: Set[str]) -> int:
    """Count the number of paragraphs that will be translated

    Args:
        doc: The Document object
        target_styles_set: Set of style names to translate

    Returns:
        Count of paragraphs to translate
    """
    count = 0
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if style_name in target_styles_set and paragraph.text.strip():
            count += 1
    return count


async def process_document_parallel(
    input_file: Path,
    output_file: Path,
    target_language: str,
    target_styles: List[str],
    openai_client: OpenAI,
    use_cache: bool = True,
    cache_dir: Path = DEFAULT_CACHE_DIR,
    max_concurrent: int = DEFAULT_MAX_CONCURRENT,
    progress_callback=None,
    cancellation_check=None,
) -> None:
    """Process a document in parallel, translating paragraphs concurrently

    Args:
        input_file: Path to the input DOCX file
        output_file: Path to save the translated DOCX file
        target_language: The target language for translation
        target_styles: List of paragraph styles to translate
        openai_client: The OpenAI client instance
        use_cache: Whether to use translation caching
        cache_dir: Directory to store cache files
        max_concurrent: Maximum number of concurrent translation requests
        progress_callback: Optional callback function for progress updates
        cancellation_check: Optional function to check if process should be cancelled
    """
    global translation_progress, total_tokens_received, total_cached_tokens

    # Reset token counters
    total_tokens_received = 0
    total_cached_tokens = 0
    translation_progress = {}

    logger.info(f"Processing document {input_file} in parallel mode")
    start_time = time.time()

    # Create async OpenAI client
    async_client = setup_async_openai_client(
        api_key=openai_client.api_key, base_url=openai_client.base_url
    )

    # Set up translation cache if enabled
    cache = None
    if use_cache:
        cache = TranslationCache(
            target_language=target_language, cache_dir=cache_dir, model=DEFAULT_MODEL
        )

    # Load the document
    doc = Document(input_file)

    # Convert target_styles to a set for faster lookups
    target_styles_set = set(target_styles)

    # Count total paragraphs to translate for progress tracking
    total_paragraphs = count_translatable_elements(doc, target_styles_set)
    logger.info(f"Found {total_paragraphs} paragraphs to translate")

    # Initial progress update
    if progress_callback:
        progress_callback(0, total_paragraphs, "Starting translation...")

    # Check if we should cancel before starting
    if cancellation_check and cancellation_check():
        logger.info("Translation cancelled before starting")
        return

    # Process paragraphs in parallel
    semaphore = asyncio.Semaphore(max_concurrent)
    tasks = []

    async def translate_with_semaphore(text, task_num, task_idx):
        """Helper function to translate text with semaphore for concurrency control"""
        task_id = f"task_{task_idx}"

        # Check if we should cancel
        if cancellation_check and cancellation_check():
            return None, task_num, task_idx

        async with semaphore:
            # Check again after acquiring semaphore
            if cancellation_check and cancellation_check():
                return None, task_num, task_idx

            translation, status = await stream_async_translation(
                text,
                target_language,
                async_client,
                cache,
                None,  # No progress bar in parallel mode
                task_id,
                cancellation_check,
            )

            # Update progress
            if progress_callback:
                processed = len(
                    [
                        t
                        for t in translation_progress.values()
                        if t.get("status") in ("completed", "error", "cached")
                    ]
                )
                progress_callback(
                    processed,
                    total_paragraphs,
                    f"Translating paragraph {processed}/{total_paragraphs}",
                )

            return translation, task_num, task_idx

    # Prepare tasks for all paragraphs that need translation
    paragraph_tasks = []
    task_idx = 0
    for i, paragraph in enumerate(doc.paragraphs):
        style_name = paragraph.style.name
        if style_name in target_styles_set and paragraph.text.strip():
            # Skip if already contains translation (for safety)
            if "\n" in paragraph.text:
                continue

            text = paragraph.text.strip()
            if text:
                # Create a task to translate this paragraph
                task = translate_with_semaphore(text, i, task_idx)
                paragraph_tasks.append((i, task_idx, task))
                task_idx += 1

    # Check if we should cancel before starting tasks
    if cancellation_check and cancellation_check():
        logger.info("Translation cancelled before tasks started")
        return

    # Start all translation tasks
    for _, task_idx, task in paragraph_tasks:
        tasks.append(asyncio.create_task(task))

    # Wait for all tasks to complete or cancellation
    translations = []
    for task in asyncio.as_completed(tasks):
        # Check for cancellation
        if cancellation_check and cancellation_check():
            # Cancel all remaining tasks
            for t in tasks:
                if not t.done():
                    t.cancel()
            logger.info("Translation cancelled during processing")
            return

        try:
            result = await task
            if result:
                translation, para_idx, task_idx = result
                translations.append((para_idx, translation))
        except asyncio.CancelledError:
            # Task was cancelled
            continue
        except Exception as e:
            logger.error(f"Error in translation task: {e}")

    # Check if we should cancel before applying translations
    if cancellation_check and cancellation_check():
        logger.info("Translation cancelled before applying translations")
        return

    # Sort translations by paragraph index
    translations.sort(key=lambda x: x[0])

    # Apply translations to paragraphs
    para_lookup = {i: para for i, para in enumerate(doc.paragraphs)}
    for para_idx, translation in translations:
        if translation:  # Skip if translation is empty (due to error or cancellation)
            add_translation_to_paragraph(para_lookup[para_idx], translation)

    # Save the document
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)

    # Final progress update
    elapsed_time = time.time() - start_time
    logger.info(f"Document processed in {elapsed_time:.2f} seconds")
    logger.info(
        f"Used {total_tokens_received} API tokens and {total_cached_tokens} cached tokens"
    )

    if progress_callback:
        progress_callback(
            total_paragraphs,
            total_paragraphs,
            f"Translation completed in {elapsed_time:.2f} seconds",
        )


def process_document(
    input_file: Path,
    output_file: Path,
    target_language: str,
    target_styles: List[str],
    openai_client: OpenAI,
    use_cache: bool = True,
    cache_dir: Path = DEFAULT_CACHE_DIR,
    parallel: bool = True,
    max_concurrent: int = DEFAULT_MAX_CONCURRENT,
    progress_callback=None,
    cancellation_check=None,
) -> None:
    """Process a document, translating specified paragraph styles

    This function wraps the parallel and sequential processing methods.

    Args:
        input_file: Path to the input DOCX file
        output_file: Path to save the translated DOCX file
        target_language: The target language for translation
        target_styles: List of paragraph styles to translate
        openai_client: The OpenAI client instance
        use_cache: Whether to use translation caching
        cache_dir: Directory to store cache files
        parallel: Whether to use parallel processing
        max_concurrent: Maximum number of concurrent translation requests
        progress_callback: Optional callback function for progress updates
        cancellation_check: Optional function to check if process should be cancelled
    """
    global translation_progress, total_tokens_received, total_cached_tokens

    # Reset token counters
    total_tokens_received = 0
    total_cached_tokens = 0
    translation_progress = {}

    if parallel:
        # Use asyncio for parallel processing
        asyncio.run(
            process_document_parallel(
                input_file,
                output_file,
                target_language,
                target_styles,
                openai_client,
                use_cache,
                cache_dir,
                max_concurrent,
                progress_callback,
                cancellation_check,
            )
        )
    else:
        # Use sequential processing
        process_document_sequential(
            input_file,
            output_file,
            target_language,
            target_styles,
            openai_client,
            use_cache,
            cache_dir,
            progress_callback,
            cancellation_check,
        )


def process_document_sequential(
    input_file: Path,
    output_file: Path,
    target_language: str,
    target_styles: List[str],
    openai_client: OpenAI,
    use_cache: bool = True,
    cache_dir: Path = DEFAULT_CACHE_DIR,
    progress_callback=None,
    cancellation_check=None,
) -> None:
    """Process a document sequentially, translating one paragraph at a time

    Args:
        input_file: Path to the input DOCX file
        output_file: Path to save the translated DOCX file
        target_language: The target language for translation
        target_styles: List of paragraph styles to translate
        openai_client: The OpenAI client instance
        use_cache: Whether to use translation caching
        cache_dir: Directory to store cache files
        progress_callback: Optional callback function for progress updates
        cancellation_check: Optional function to check if process should be cancelled
    """
    global translation_progress, total_tokens_received, total_cached_tokens

    logger.info(f"Processing document {input_file} in sequential mode")
    start_time = time.time()

    # Set up translation cache if enabled
    cache = None
    if use_cache:
        cache = TranslationCache(
            target_language=target_language, cache_dir=cache_dir, model=DEFAULT_MODEL
        )

    # Load the document
    doc = Document(input_file)

    # Convert target_styles to a set for faster lookups
    target_styles_set = set(target_styles)

    # Count total paragraphs to translate for progress tracking
    total_paragraphs = count_translatable_elements(doc, target_styles_set)
    logger.info(f"Found {total_paragraphs} paragraphs to translate")

    # Use tqdm for progress tracking in CLI mode
    progress_bar = tqdm(total=total_paragraphs, desc="Translating paragraphs")

    # Track current progress
    processed_paragraphs = 0

    # Process each paragraph
    for paragraph in doc.paragraphs:
        # Check if we should cancel
        if cancellation_check and cancellation_check():
            logger.info("Translation cancelled")
            progress_bar.close()
            return

        style_name = paragraph.style.name
        if style_name in target_styles_set and paragraph.text.strip():
            # Skip if already contains translation (for safety)
            if "\n" in paragraph.text:
                continue

            text = paragraph.text.strip()
            if text:
                # Translate the paragraph
                task_id = f"para_{processed_paragraphs}"
                translated_text = stream_translation(
                    text, target_language, openai_client, cache, progress_bar, task_id
                )

                # Add translation to the paragraph
                if translated_text:
                    add_translation_to_paragraph(paragraph, translated_text)

                # Update progress
                processed_paragraphs += 1
                if progress_callback:
                    progress_callback(
                        processed_paragraphs,
                        total_paragraphs,
                        f"Translating paragraph {processed_paragraphs}/{total_paragraphs}",
                    )

    # Save the document
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)

    # Close progress bar
    progress_bar.close()

    # Final progress update
    elapsed_time = time.time() - start_time
    logger.info(f"Document processed in {elapsed_time:.2f} seconds")
    logger.info(
        f"Used {total_tokens_received} API tokens and {total_cached_tokens} cached tokens"
    )

    if progress_callback:
        progress_callback(
            total_paragraphs,
            total_paragraphs,
            f"Translation completed in {elapsed_time:.2f} seconds",
        )


def clear_translation_caches(cache_dir: Path = DEFAULT_CACHE_DIR) -> None:
    """Clear all translation cache files

    Args:
        cache_dir: Directory where cache files are stored
    """
    logger.info(f"Clearing translation caches in {cache_dir}")

    if not cache_dir.exists():
        logger.info(f"Cache directory {cache_dir} does not exist. Nothing to clear.")
        return

    # Find and remove all cache files
    for cache_file in cache_dir.glob("cache_*.json"):
        try:
            cache_file.unlink()
            logger.info(f"Removed cache file: {cache_file.name}")
        except Exception as e:
            logger.error(f"Failed to remove cache file {cache_file.name}: {e}")

    logger.info("Translation caches cleared")
