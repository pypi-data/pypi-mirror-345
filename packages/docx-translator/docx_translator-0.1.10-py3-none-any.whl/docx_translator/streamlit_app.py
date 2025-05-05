import os
import tempfile
import argparse
import threading
from queue import Queue
from pathlib import Path
import streamlit as st
from docx import Document
from dotenv import load_dotenv
import time
from loguru import logger
import logging

for name, l in logging.root.manager.loggerDict.items():
    if "streamlit" in name:
        l.disabled = True

# Load environment variables from .env file
load_dotenv()

# Import from docx_translator.translator package
from docx_translator.translator import (
    setup_openai_client,
    process_document,
    DEFAULT_MODEL,
    DEFAULT_CACHE_DIR,
    DEFAULT_MAX_CONCURRENT,
)

# Parse command line arguments
parser = argparse.ArgumentParser(description="Streamlit app for DocxTranslator")
parser.add_argument(
    "--cache-dir", type=str, default=str(DEFAULT_CACHE_DIR), help="Cache directory"
)
parser.add_argument(
    "--model", type=str, default=DEFAULT_MODEL, help="OpenAI model to use"
)
parser.add_argument(
    "--target-language", type=str, default="Spanish", help="Default target language"
)
args, _ = parser.parse_known_args()
cache_dir = Path(args.cache_dir)
default_model = args.model
default_target_language = args.target_language

st.set_page_config(
    page_title="DocxTranslator",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Initialize session state for translation control
if "translation_stopped" not in st.session_state:
    st.session_state.translation_stopped = False
if "translation_running" not in st.session_state:
    st.session_state.translation_running = False
if "tokens_info" not in st.session_state:
    st.session_state.tokens_info = {"api": 0, "cached": 0}
if "start_translation" not in st.session_state:
    st.session_state.start_translation = False
if "translation_params" not in st.session_state:
    st.session_state.translation_params = None
if "translation_thread" not in st.session_state:
    st.session_state.translation_thread = None
if "translation_queue" not in st.session_state:
    st.session_state.translation_queue = Queue()
if "result" not in st.session_state:
    st.session_state.result = None
if "last_rerun_time" not in st.session_state:
    st.session_state.last_rerun_time = time.time()
if "rerun_interval" not in st.session_state:
    st.session_state.rerun_interval = (
        0.2  # Reduce to 0.2 seconds for more frequent updates
    )
# Always initialize progress tracking values
if "progress_value" not in st.session_state:
    st.session_state.progress_value = 0
if "progress_description" not in st.session_state:
    st.session_state.progress_description = "Initializing translation..."
if "token_text" not in st.session_state:
    st.session_state.token_text = "Tokens: 0 API + 0 cached = 0 total"
# Add tracking for ETA calculation
if "progress_start_time" not in st.session_state:
    st.session_state.progress_start_time = None
if "last_progress_value" not in st.session_state:
    st.session_state.last_progress_value = 0
if "last_progress_time" not in st.session_state:
    st.session_state.last_progress_time = None
# Add tracking for the current ETA (initialized to empty string rather than None)
if "current_eta" not in st.session_state:
    st.session_state["current_eta"] = ""
if "download_available" not in st.session_state:
    st.session_state.download_available = False

# Sidebar for API key and settings
with st.sidebar:
    st.title("DocxTranslator Settings")

    # API key input - pre-fill from environment variable if available
    api_key = st.text_input(
        "OpenAI API Key",
        value=os.environ.get("OPENAI_API_KEY", ""),
        type="password",
        help="Your OpenAI API key. This will not be stored.",
    )

    # Base URL input for custom endpoints - pre-fill from environment variable if available
    base_url = st.text_input(
        "OpenAI API Base URL (Optional)",
        value=os.environ.get("OPENAI_BASE_URL", ""),
        help="Custom OpenAI API base URL (for proxies or custom endpoints)",
    )
    if not base_url:
        base_url = None

    # Model input - pre-fill with default model
    model = st.text_input(
        "OpenAI Model",
        value=os.environ.get("OPENAI_MODEL", default_model),
        help="The OpenAI model to use for translation (e.g., gpt-4-turbo, gpt-4o, gpt-3.5-turbo)",
    )

    # Target language
    target_language = st.text_input(
        "Target Language",
        value=os.environ.get("TARGET_LANGUAGE", default_target_language),
        help="The language to translate to (e.g., Spanish, French, Japanese)",
    )

    # Advanced settings in expander
    with st.expander("Advanced Settings"):
        use_cache = st.toggle(
            "Use Translation Cache",
            value=True,
            help="Cache translations to avoid redundant API calls",
        )

        clear_cache = st.checkbox(
            "Clear Cache Before Starting",
            value=False,
            help="Clear the translation cache for this language before starting translation",
        )

        parallel = st.toggle(
            "Use Parallel Processing",
            value=True,
            help="Process translations in parallel for faster results",
        )

        max_concurrent = st.slider(
            "Max Concurrent Requests",
            min_value=1,
            max_value=20,
            value=DEFAULT_MAX_CONCURRENT,
            help="Maximum number of concurrent translation requests",
        )

    # About section
    with st.expander("About DocxTranslator"):
        st.write(
            """
            DocxTranslator uses OpenAI's API to translate Word documents.
            The application preserves formatting and document structure.
            
            Translations are added below the original text in each paragraph.
            
            For command line usage, check out the [GitHub repository](https://github.com/john-theo/docx-translator.git).
            """
        )

# Main content
st.title("📄 DocxTranslator")
st.write("Upload a Word document (.docx) to translate using OpenAI API")

# File uploader
uploaded_file = st.file_uploader("Choose a Word document", type=["docx"])


# Define a function to reset translation state properly
def reset_translation_state():
    """Reset all translation state properly to allow starting a new translation"""
    # Reset flags
    st.session_state.translation_running = False
    st.session_state.translation_stopped = False

    # Clear the queue
    while not st.session_state.translation_queue.empty():
        try:
            st.session_state.translation_queue.get_nowait()
        except:
            break

    # Reset progress display for next run
    st.session_state.progress_value = 0
    st.session_state.progress_description = "Initializing translation..."
    st.session_state.token_text = "Tokens: 0 API + 0 cached = 0 total"
    # Reset ETA tracking but preserve ETA for display
    st.session_state.progress_start_time = None
    st.session_state.last_progress_value = 0
    st.session_state.last_progress_time = None
    # Clear the stored ETA
    st.session_state["current_eta"] = ""

    # Clear UI elements - except result_area if we have a download available
    try:
        # Only try to clear these if they exist in the current context
        if "status_text" in globals():
            status_text.empty()
        if "progress_bar" in globals():
            progress_bar.empty()
        # Don't clear result_area if we have a download available
        if "result_area" in globals() and not st.session_state.get(
            "download_available", False
        ):
            result_area.empty()
    except:
        pass  # Ignore errors if elements don't exist in context

    # Force UI update
    return True  # Return True to indicate queue should be updated


# Define a function to clear download state
def clear_download_state():
    """Clear the download state before starting a new translation"""
    st.session_state.download_available = False
    st.session_state.result = None
    if "result_area" in globals():
        result_area.empty()


# Translation thread function
def translation_thread_func(params, queue):
    try:
        # Set OpenAI API key
        os.environ["OPENAI_API_KEY"] = params["api_key"]

        # Set up the OpenAI client
        client = setup_openai_client(params["api_key"], params["base_url"])

        # Set the global model if different from default
        if params["model"] != DEFAULT_MODEL:
            from docx_translator import translator

            translator.DEFAULT_MODEL = params["model"]

        # Clear the cache if requested
        if params["clear_cache"] and params["use_cache"]:
            try:
                cache_file = (
                    Path(params["cache_dir"])
                    / f"cache_{params['target_language'].lower().replace(' ', '_')}_{params['model'].replace('-', '_')}.json"
                )
                if cache_file.exists():
                    os.remove(cache_file)
                    queue.put(
                        (
                            "info",
                            f"Cleared translation cache for {params['target_language']}",
                        )
                    )
            except Exception as e:
                queue.put(("warning", f"Failed to clear cache: {str(e)}"))

        # Create a cancellation checker function
        stop_requested = False

        def check_cancellation():
            nonlocal stop_requested
            # If already flagged for stopping, return True immediately
            if stop_requested:
                return True

            # Check if there's a stop command in the queue without blocking other messages
            try:
                if not queue.empty():
                    cmd = queue.get_nowait()
                    if cmd == "stop":
                        stop_requested = True
                        queue.put("stop")  # Put it back for other parts of the code
                        # Let other parts of code know cancellation was requested
                        logger.warning("User requested cancellation of translation")
                        return True
                    else:
                        # Put back any other messages in a way that doesn't block progress updates
                        queue.put(cmd)
            except Exception:
                pass  # Ignore queue checking errors

            return stop_requested

        # Start a heartbeat thread to ensure UI updates happen even when process is busy
        def heartbeat_function():
            # Only run while translation is active and not stopped
            while not stop_requested and threading.current_thread().is_alive():
                try:
                    # Check if we should continue running
                    if check_cancellation():
                        break

                    # Send heartbeat message to keep UI updating - but don't modify the progress value
                    # Just send a message to keep the queue active without changing progress
                    queue.put(("heartbeat", None))
                    time.sleep(0.2)  # Send a heartbeat every 200ms
                except Exception as e:
                    # Just log and continue - don't let errors in heartbeat stop translation
                    print(f"Heartbeat error: {str(e)}")
                    time.sleep(1)  # Wait longer on error

        # Start heartbeat in a daemon thread so it auto-terminates when main thread exits
        heartbeat_thread = threading.Thread(target=heartbeat_function, daemon=True)
        heartbeat_thread.start()

        # Define progress callback function
        def update_progress(current, total, description):
            # Check if translation should be stopped
            if check_cancellation():
                # Signal cancellation without raising exception (avoids traceback)
                queue.put(("progress", (0, "Cancellation requested...")))
                return False  # Return value indicates cancellation

            # Update progress through queue with normalized progress (0-1)
            progress = min(current / total, 1.0) if total > 0 else 0

            # Get token counts from global variables in translator module
            from docx_translator import translator

            api_tokens = translator.total_tokens_received
            cached_tokens = translator.total_cached_tokens

            # Show more detailed progress with current/total, percentage
            # Check if we have a previously calculated ETA stored in session state
            eta_part = ""
            # Use get() method to safely access the variable with a default value
            if st.session_state.get("current_eta") is not None:
                eta_part = st.session_state.get("current_eta")

            # Always create description with the most recent ETA
            # Basic progress description
            simple_description = f"Progress: {current}/{total}, {int(progress * 100)}%"

            # Always add the ETA if available
            if eta_part:
                simple_description += eta_part

            # Always push updates to the queue, even for small changes
            queue.put(("progress", (progress, simple_description)))
            queue.put(("tokens", {"api": api_tokens, "cached": cached_tokens}))

            # Return True to indicate successful update
            return True

        try:
            # Process the document with progress updates and cancellation checking
            process_document(
                Path(params["input_path"]),
                Path(params["output_path"]),
                params["target_language"],
                params["styles_to_translate"],
                client,
                use_cache=params["use_cache"],
                cache_dir=Path(params["cache_dir"]),
                parallel=params["parallel"],
                max_concurrent=params["max_concurrent"],
                progress_callback=update_progress,
                cancellation_check=check_cancellation,
            )

            # Only proceed with file reading if we haven't been cancelled
            if not check_cancellation():
                # Read output file
                with open(params["output_path"], "rb") as output_file:
                    output_data = output_file.read()

                # Send success message and file data
                queue.put(
                    (
                        "success",
                        {
                            "file_data": output_data,
                            "filename": f"translated_{params['uploaded_filename']}",
                        },
                    )
                )
            else:
                queue.put(("interrupted", "Translation process was stopped by user."))

        except InterruptedError as e:
            # Handle cancellation without propagating the exception to avoid traceback
            logger.warning(f"Translation was stopped by user: {str(e)}")
            queue.put(("interrupted", "Translation process was stopped by user."))
        except Exception as e:
            queue.put(("error", f"Error during translation: {str(e)}"))

        # Clean up temporary files
        try:
            if Path(params["input_path"]).exists():
                os.unlink(params["input_path"])
            if Path(params["output_path"]).exists():
                os.unlink(params["output_path"])
        except:
            pass  # Ignore deletion errors

        # Remove API key from environment
        if "OPENAI_API_KEY" in os.environ:
            del os.environ["OPENAI_API_KEY"]

    except Exception as e:
        queue.put(("error", f"Error setting up translation: {str(e)}"))
        if "OPENAI_API_KEY" in os.environ:
            del os.environ["OPENAI_API_KEY"]

    finally:
        # Signal that thread is done
        queue.put(("done", None))


if uploaded_file:
    # Display document info
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_input_path = Path(tmp_file.name)

    # Get document details
    doc = Document(tmp_input_path)

    # Extract styles from document
    available_styles = []
    for para in doc.paragraphs:
        if para.style.name and para.style.name not in available_styles:
            available_styles.append(para.style.name)

    # If tables exist, also check table styles
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.style.name and para.style.name not in available_styles:
                        available_styles.append(para.style.name)

    # Sort styles alphabetically but keep "Normal" first if present
    if "Normal" in available_styles:
        available_styles.remove("Normal")
        available_styles = ["Normal"] + sorted(available_styles)
    else:
        available_styles.sort()

    # Let user select styles to translate
    styles_to_translate = st.multiselect(
        "Select styles to translate",
        options=available_styles,
        default=["Normal"] if "Normal" in available_styles else available_styles[:1],
        help="Only text with these styles will be translated",
    )

    # Count paragraphs with selected styles
    paragraph_count = 0
    for para in doc.paragraphs:
        if para.style.name in styles_to_translate and para.text.strip():
            paragraph_count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.style.name in styles_to_translate and para.text.strip():
                        paragraph_count += 1

    st.write(f"Document has {paragraph_count} paragraphs with selected styles.")

    # Warning about tokens/cost
    if paragraph_count > 0:
        st.warning(
            f"Translating {paragraph_count} paragraphs may consume a significant number of tokens. "
            f"Make sure you have sufficient OpenAI API credits."
        )

    # Translate button
    col1, col2, col3 = st.columns([4, 1, 1])

    # Create initial processing elements
    progress_placeholder = st.empty()
    progress_bar = st.empty()
    status_text = st.empty()
    result_area = st.empty()

    # Prepare the translation function
    def set_translation_params():
        # Only prepare translation if all conditions are met
        if not api_key or not styles_to_translate or paragraph_count == 0:
            return

        # Clear any previous download state
        clear_download_state()

        # Reset the flags
        st.session_state.translation_stopped = False
        st.session_state.tokens_info = {"api": 0, "cached": 0}
        st.session_state.result = None

        # Clear the queue to remove any previous messages
        while not st.session_state.translation_queue.empty():
            st.session_state.translation_queue.get()

        # Set translation running flag
        st.session_state.translation_running = True

        # Prepare parameters for translation
        with tempfile.NamedTemporaryFile(
            suffix=".docx", delete=False
        ) as tmp_output_file:
            tmp_output_path = Path(tmp_output_file.name)

        st.session_state.translation_params = {
            "input_path": str(tmp_input_path),
            "output_path": str(tmp_output_path),
            "target_language": target_language,
            "styles_to_translate": styles_to_translate,
            "api_key": api_key,
            "base_url": base_url,
            "model": model,
            "use_cache": use_cache,
            "cache_dir": str(cache_dir),
            "parallel": parallel,
            "max_concurrent": max_concurrent,
            "clear_cache": clear_cache,
            "uploaded_filename": uploaded_file.name,
        }

        # Signal to start translation on next rerun
        st.session_state.start_translation = True

        # Force a rerun to update the UI
        st.rerun()

    translate_button = col1.button(
        "Translate Document",
        type="primary",
        disabled=paragraph_count == 0
        or not api_key
        or not styles_to_translate
        or st.session_state.translation_running,
        on_click=set_translation_params,
    )

    # Stop button to cancel the translation process
    def stop_translation():
        if st.session_state.translation_running:
            # Put stop command in the queue for the worker thread to pick up
            st.session_state.translation_queue.put("stop")

            # Set the flag to indicate cancellation was requested
            st.session_state.translation_stopped = True

            # Update the status text to show cancellation is in progress
            st.warning(
                "Stopping translation... Please wait for running tasks to complete."
            )

            # Force a rerun to update the UI
            st.rerun()

    stop_button = col2.button(
        "Stop",
        type="secondary",
        disabled=not st.session_state.translation_running,
        on_click=stop_translation,
    )

    # If we should start translation (triggered by the previous run)
    if st.session_state.start_translation and st.session_state.translation_params:
        # Reset the flag to prevent restarting translation
        st.session_state.start_translation = False

        # Initialize the progress bar and status
        progress_bar.progress(0)
        status_text.text("Initializing translation...")

        # Reset ETA tracking for the new translation
        st.session_state.progress_start_time = time.time()
        st.session_state.last_progress_time = time.time()
        st.session_state.last_progress_value = 0

        # Start the translation thread
        st.session_state.translation_thread = threading.Thread(
            target=translation_thread_func,
            args=(
                st.session_state.translation_params,
                st.session_state.translation_queue,
            ),
            daemon=True,
        )
        st.session_state.translation_thread.start()

    # Check for messages from the translation thread
    if st.session_state.translation_running:
        # Always display current status using session state values
        progress_bar.progress(st.session_state.progress_value)

        # Always show the progress description - this ensures it stays visible
        current_description = st.session_state.progress_description

        # Make sure current description always includes the ETA if available
        if " (ETA:" not in current_description and st.session_state.get("current_eta"):
            current_description = current_description.split(" (ETA:")[
                0
            ] + st.session_state.get("current_eta")

        # Always update the status text with current progress
        status_text.text(current_description)

        # Update token information
        progress_placeholder.text(st.session_state.token_text)

        # Process messages from the queue
        queue_updated = False
        queued_messages = []

        # First collect all messages without blocking
        while not st.session_state.translation_queue.empty():
            try:
                message = st.session_state.translation_queue.get_nowait()
                # Validate message format
                if message == "stop":
                    queued_messages.append(("stop", None))
                elif isinstance(message, tuple) and len(message) == 2:
                    queued_messages.append(message)
                else:
                    status_text.text(f"Received invalid message format: {message}")
                queue_updated = True
            except Exception as e:
                status_text.text(f"Error getting message from queue: {str(e)}")
                break

        # Process all collected messages
        if queued_messages:
            # Sort messages by type to ensure progress updates are processed in the right order
            # Process 'progress' messages last to ensure we have the latest progress
            progress_messages = []
            token_messages = []
            other_messages = []
            heartbeat_messages = []

            for message in queued_messages:
                try:
                    message_type, message_data = message
                    if message_type == "progress":
                        progress_messages.append(message)
                    elif message_type == "tokens":
                        token_messages.append(message)
                    elif message_type == "heartbeat":
                        heartbeat_messages.append(message)
                    else:
                        other_messages.append(message)
                except Exception as e:
                    status_text.text(f"Error unpacking message: {str(e)}")

            # Process in this order: other messages, token updates, progress updates, heartbeats last
            for message in (
                other_messages + token_messages + progress_messages + heartbeat_messages
            ):
                try:
                    message_type, message_data = message

                    if message_type == "progress":
                        progress_value, description = message_data
                        # Only update the progress if it's increasing or equal
                        # This prevents the progress bar from bouncing backwards
                        if progress_value >= st.session_state.progress_value:
                            current_time = time.time()

                            # Initialize start time if not set
                            if st.session_state.progress_start_time is None:
                                st.session_state.progress_start_time = current_time
                                st.session_state.last_progress_time = current_time
                                st.session_state.last_progress_value = 0

                            # Calculate ETA based on progress rate
                            if (
                                progress_value > st.session_state.last_progress_value
                                and current_time > st.session_state.last_progress_time
                            ):
                                # Calculate progress rate (progress per second)
                                elapsed = (
                                    current_time - st.session_state.progress_start_time
                                )
                                if elapsed > 0:
                                    # Calculate progress per second - ensure it's positive
                                    progress_rate = max(
                                        0.0001, progress_value / elapsed
                                    )

                                    # Calculate remaining time based on rate
                                    remaining_progress = 1.0 - progress_value
                                    eta_seconds = remaining_progress / progress_rate
                                    logger.debug(
                                        f"Progress calculation: rate={progress_rate}, remaining={remaining_progress}, eta={eta_seconds}"
                                    )

                                    # Format minutes and seconds
                                    if eta_seconds > 0:
                                        eta_minutes = int(eta_seconds // 60)
                                        eta_secs = int(eta_seconds % 60)
                                        eta_str = (
                                            f" (ETA: {eta_minutes:02d}:{eta_secs:02d})"
                                        )

                                        # Store the ETA in session state for reuse
                                        st.session_state["current_eta"] = eta_str

                                        # Make sure to split by ETA to remove any previous ETA information
                                        description = (
                                            description.split(" (ETA:")[0] + eta_str
                                        )
                                        logger.debug(f"New ETA calculated: {eta_str}")

                            # Update last progress for next calculation
                            st.session_state.last_progress_value = progress_value
                            st.session_state.last_progress_time = current_time

                            # Update the UI
                            st.session_state.progress_value = progress_value

                            # Make sure we always include the ETA
                            if " (ETA:" not in description and st.session_state.get(
                                "current_eta"
                            ):
                                description = description.split(" (ETA:")[
                                    0
                                ] + st.session_state.get("current_eta")

                            # Store in session state for the main UI loop to pick up
                            st.session_state.progress_description = description

                            # Debug output for tracking
                            logger.debug(f"Progress update in handler: {description}")

                            # No need to update UI here - the main UI loop will handle this consistently

                    elif message_type == "tokens":
                        st.session_state.tokens_info = message_data
                        st.session_state.token_text = f"Tokens: {message_data['api']} API + {message_data['cached']} cached = {message_data['api'] + message_data['cached']} total"

                    elif message_type == "info":
                        st.info(message_data)

                    elif message_type == "warning":
                        st.warning(message_data)

                    elif message_type == "error":
                        # Don't clear status_text to maintain progress display
                        # status_text.empty()

                        # Show error
                        result_area.error(message_data)

                        # Reset all translation state on error
                        queue_updated = reset_translation_state()

                        # Force a rerun to update button states
                        st.rerun()

                    elif message_type == "success":
                        # Only clear status_text at the end of a successful job
                        status_text.empty()

                        # Set the result and download availability flag
                        st.session_state.result = message_data
                        st.session_state.download_available = True

                        # Show completion message
                        result_area.success("✅ Translation completed!")

                        # Display final token summary
                        api_tokens = st.session_state.tokens_info["api"]
                        cached_tokens = st.session_state.tokens_info["cached"]
                        total_tokens = api_tokens + cached_tokens
                        result_area.info(
                            f"📊 Token usage: {api_tokens} API + {cached_tokens} cached = {total_tokens} total"
                        )

                        # Show download button for the result
                        result_area.download_button(
                            label="Download Translated Document",
                            data=message_data["file_data"],
                            file_name=message_data["filename"],
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_button",
                            type="primary",
                        )

                        # Reset translation state but preserve result data
                        queue_updated = reset_translation_state()

                        # No rerun to ensure download button persists

                    elif message_type == "interrupted":
                        # Only clear status_text at the end of a job
                        status_text.empty()

                        # Show the interruption message
                        result_area.warning(message_data)

                        # Display token summary even if interrupted
                        api_tokens = st.session_state.tokens_info["api"]
                        cached_tokens = st.session_state.tokens_info["cached"]
                        total_tokens = api_tokens + cached_tokens
                        result_area.info(
                            f"📊 Token usage before interruption: {api_tokens} API + {cached_tokens} cached = {total_tokens} total"
                        )

                        # Reset translation state to allow starting a new translation
                        queue_updated = reset_translation_state()

                        # Force an immediate rerun to update button states
                        st.rerun()

                    elif message_type == "done":
                        # Thread is finished, clean up
                        st.session_state.translation_thread = None

                        # If translation was interrupted or has any flags still set, clean up completely
                        if (
                            st.session_state.translation_running
                            or st.session_state.translation_stopped
                        ):
                            # Clear all UI elements at the very end
                            status_text.empty()
                            progress_bar.empty()
                            progress_placeholder.empty()

                            # Always ensure UI state is reset properly
                            queue_updated = reset_translation_state()

                            # Force an immediate rerun to update button states - MUST BE LAST
                            st.rerun()

                    elif message_type == "stop":
                        # Put the stop command back in the queue for the translation thread to see
                        st.session_state.translation_queue.put("stop")
                        st.warning("Stopping translation process... Please wait.")

                    elif message_type == "heartbeat":
                        # Heartbeat message just forces a rerun without changing any UI state
                        # This keeps the UI responsive during long-running operations
                        pass

                except Exception as e:
                    status_text.text(f"Error processing message: {str(e)}")

        # Check if it's time for a periodic rerun
        current_time = time.time()
        time_since_last_rerun = current_time - st.session_state.last_rerun_time

        # Always force a rerun while translation is running
        if st.session_state.translation_running:
            # Update the last rerun time
            st.session_state.last_rerun_time = current_time

            # Add a small sleep to prevent too rapid reruns that might overload the browser
            time.sleep(0.1)

            # Force a rerun to update the UI
            st.rerun()

else:
    st.info("Please upload a Word document (.docx) to begin translation.")
