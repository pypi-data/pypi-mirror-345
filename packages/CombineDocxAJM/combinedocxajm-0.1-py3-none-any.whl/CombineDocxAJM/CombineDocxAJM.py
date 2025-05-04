"""
# Module: CombineDocx.py

This module provides functionality for combining multiple 'docx' files into a single file, particularly focusing on the
`CombineDocx` class.
The primary class and methods include:

## Classes:
- **FileSorter**: Handles sorting of files, particularly by extracting and sorting by numerical suffixes in file names.
  - **Attributes**:
    - `NEGATIVE_INFINITY`: Constant used to represent negative infinity.
    - `directory`: Path to the directory containing files to be sorted.
  - **Methods**:
    - `__init__`: Initializes the `FileSorter` with directory and logger.
    - `_extract_last_number`: Extracts the last number from a file name.
    - `sort_by_last_number`: Sorts files based on their numerical suffixes.

- **CombineDocx**: Combines multiple 'docx' files into a single master file.
  - **Attributes**:
    - `_save_path`: Path where the combined file is saved.
    - `auto_overwrite_combined_file`: Boolean indicating whether to auto-overwrite the combined file.
    - `file_list`: List of files to be combined.
    - `_logger`: Logger instance for logging events.
    - `directory_to_combine`: Directory where the combined file will be saved.
    - `combined_filename`: Name of the combined file.
    - `master_filename`: Path to the master file.
    - `_config`: Configuration object for retrieving settings.
  - **Methods**:
    - `__init__`: Initializes the `CombineDocx` class with necessary attributes.
    - `_get_config_auto_overwrite`: Retrieves the auto-overwrite setting from the configuration.
    - `file_age`: Calculates the age of a file based on its creation or modification time.
    - `save_path`: Property that returns the save path for the combined file.
    - `_are_you_sure`: Static method that asks for confirmation of an action.
    - `combine_all_docx`: Combines all 'docx' files into the master file.

- **CombineSortedDocx**: Extends the functionality of `CombineDocx` to combine files that have been sorted.
  - **Attributes**:
    - `file_list`: List of files to be combined.
    - `_logger`: Logger instance for logging events.
    - `directory_to_combine`: Directory where the combined file will be saved.
  - **Methods**:
    - `__init__`: Initializes the `CombineSortedDocx` class with necessary attributes.

## Usage:
The `CombineDocx` class can be instantiated and used to combine multiple 'docx' files based on user-defined settings
or default configurations. Additional helper methods provide functionality
for handling file properties and user confirmations.

Example:
```python
from CombineDocx import CombineDocx
combiner = CombineDocx(master_filename="master.docx", file_list=["file1.docx", "file2.docx"], directory_to_combine="docs")
combined_file_path = combiner.combine_all_docx()
print(f"Combined file saved at: {combined_file_path}")
```

This example demonstrates how to create an instance of the `CombineDocx` class and use it to combine
multiple 'docx' files into a single master file.
"""
import datetime
from logging import Logger, getLogger, basicConfig
from pathlib import Path
from re import findall

import questionary
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docxcompose.composer import Composer


class FileSorter:
    """
    Class FileSorter:

    A class for sorting files in a directory based on their last number in the filename.

    Attributes:
    - directory (Path): The directory to sort.

    Methods:
    - __init__(directory_to_sort: str or Path): Initializes a FileSorter instance with the specified directory to sort.
    - sort_by_last_number(): Returns a list of files in the directory, sorted by the last number in the filename.

    Static Methods:
    - _extract_last_number(filename): Extracts the last number from a filename.

    """
    NEGATIVE_INFINITY = float('-inf')

    def __init__(self, directory_to_sort: str or Path, **kwargs):
        self._logger = kwargs.get('logger', Logger("dummy_logger"))
        self.directory = Path(directory_to_sort)

    def _extract_last_number(self, filename):
        """
        Extract the last number from a filename.
        Args:
            filename (str): The name of the file.
        Returns:
            int: The last number found in the filename or negative infinity if no numbers are found.
        """
        numbers = findall(r'\d+', filename)
        if not numbers:
            self._logger.debug(f"no numbers found in {filename}")
        else:
            self._logger.debug(f"numbers found in {filename}: {numbers}")
        return int(numbers[-1]) if numbers else FileSorter.NEGATIVE_INFINITY

    def sort_by_last_number(self):
        """
        Returns a list of files in the directory, sorted by the last number in the filename.
        Returns:
            list: List of filenames ordered based on the last number in the filenames.
        """
        try:
            files = [str(Path(self.directory, f.name).as_posix()) for f in self.directory.iterdir() if f.is_file()]
        except FileNotFoundError as e:
            self._logger.error(e, exc_info=True)
            raise e

        return sorted(files, key=self._extract_last_number)


class CombineDocx:
    """
    This module provides a class, `CombineDocx`, that can be used to combine multiple .docx files into a single .docx file.

    Class:
        - `CombineDocx`

    Methods:
        - `__init__(self, master_filename, file_list, directory_to_combine, **kwargs)`
        - `_get_config_auto_overwrite(self)`
        - `file_age(self, file_path, **kwargs)`
        - `save_path(self)`
        - `_are_you_sure(action)`
        - `combine_all_docx(self)`

    ---

    `CombineDocx` class:

        The `CombineDocx` class is used to combine multiple .docx files into a single .docx file.

        Methods:

        - `__init__(self, master_filename, file_list, directory_to_combine, **kwargs)`

            This method is the initializer of the `CombineDocx` class. It takes the following parameters:

            - `master_filename` (str): The filename of the master .docx file.
            - `file_list` (list): A list of filenames of the .docx files to be combined.
            - `directory_to_combine` (str): The directory where the combined .docx file will be saved.
            - `**kwargs`: Additional keyword arguments.

            Keyword Arguments:
            - `logger` (Logger, optional): An instance of the Logger class for logging information. Defaults to a dummy logger.
            - `config` (ConfigParser, optional): An instance of the ConfigParser class for accessing configuration values. Defaults to None.
            - `combined_filename` (str, optional): The filename of the combined .docx file. Defaults to "{directory_to_combine}_combined_file.docx".
            - `auto_overwrite_combined_file` (bool, optional): If True, the combined file will be automatically overwritten if it already exists. Defaults to the value specified in the configuration file.

    ---

    `CombineDocx` methods:

        - `_get_config_auto_overwrite(self)`

            This method is used to get the value of the `auto_overwrite_combined_file` from the configuration file, if it exists.

            Returns:
            - None

        - `file_age(self, file_path, **kwargs)`

            This method calculates the age of a file based on its creation time or modification time.

            Parameters:
            - `file_path` (pathlib.Path): The path to the file.
            - `use_create_time` (bool, optional): If True, the file creation time will be used to calculate the age. Defaults to False.
            - `use_modify_time` (bool, optional): If True, the file modification time will be used to calculate the age. Defaults to True.

            Returns:
            - `datetime.timedelta`: The age of the file as a timedelta object.

            Raises:
            - `AttributeError`: If both `use_create_time` and `use_modify_time` are True or if both are False.

        - `@property save_path(self)`

            This property is used to access the path where the file is to be saved.

            Returns:
            - `pathlib.Path`: The path of the file to be saved.

            Raises:
            - `FileExistsError`: If the save path already exists and the user chooses not to overwrite it.

        - `@staticmethod _are_you_sure(action)`

            This static method asks the user if they are sure about performing a specific action.

            Parameters:
            - `action` (str): The action to be performed.

            Returns:
            - `bool`: True if the user is sure, False otherwise.

        - `combine_all_docx(self)`

            This method combines multiple .docx files into a single .docx file.

            Returns:
            - `str`: The path of the combined .docx file.

    ---

    Example:
    ```
    # Create an instance of the CombineDocx class
    combiner = CombineDocx("example-master.docx", ["example1.docx", "example2.docx"], "/path/to/directory")

    # Combine the .docx files
    combined_file_path = combiner.combine_all_docx()

    print(f"Combined file saved to: {combined_file_path}")
    ```

    ---

    Note: This module assumes that the necessary libraries (`python-docx` and `questionary`) are imported properly before calling the methods.
    """

    def __init__(self, master_filename, file_list, directory_to_combine, **kwargs):
        self._logger = kwargs.get('logger', getLogger('CombineDocx'))
        self._check_fallback_logger_config()
        self._config = kwargs.get('config', None)

        self.master_filename = Path(master_filename).as_posix()
        self.master_document = Document(str(self.master_filename))
        self.composer = Composer(self.master_document)

        self._file_list = []
        self.file_list = file_list
        # TODO: should this be renamed save_dir?
        self.directory_to_combine = Path(directory_to_combine)

        self.combined_filename = kwargs.get('combined_filename',
                                            Path(f'{self.directory_to_combine.name}_combined_file.docx'))

        # declaring auto_overwrite_combined_file like this, allows for the kwarg to supersede the config value
        self.auto_overwrite_combined_file = None
        self._get_config_auto_overwrite()
        if kwargs.get('auto_overwrite_combined_file') is not None:
            self._logger.info(f"kwarg auto_overwrite_combined_file detected")
        self.auto_overwrite_combined_file = bool(kwargs.get('auto_overwrite_combined_file',
                                                            self.auto_overwrite_combined_file))

        self.save_path = kwargs.get('save_path', Path(self.directory_to_combine, self.combined_filename))

    def _check_fallback_logger_config(self, default_logger_name='CombineDocx'):
        if self._logger.name == default_logger_name:
            basicConfig(level='INFO')
            self._logger.info('using basic config')

    def get_paragraph_by_content(self, section_content):
        section_index, paragraph_index = self._get_section_and_relative_paragraph_index(section_content=section_content)
        return self._get_paragraph_by_relative_index(section_index, paragraph_index)

    def _get_section_and_relative_paragraph_index(self, section_content, **kwargs):
        def _compare_text(strictness, search_content, candidate_content):
            if strictness == 'exact':
                return candidate_content == search_content
            elif strictness == 'partial':
                return search_content in candidate_content
            else:
                raise ValueError(f"Invalid strictness value: {strictness}, expected 'exact' or 'partial'")

        search_strictness = kwargs.get('search_strictness', 'exact')
        min_section_index = kwargs.get('min_section_index', 4)

        for section_enumeration, s in enumerate(self.master_document.sections):
            for paragraph_enumeration, content in enumerate(s.iter_inner_content()):
                if (hasattr(content, 'text')
                        and _compare_text(search_strictness, section_content, content.text)
                        and section_enumeration > min_section_index):
                    print(f'section with \'{section_content}\' found in document\n'
                          f'section # {section_enumeration}\n'
                          f'Paragraph # {paragraph_enumeration}')
                    return section_enumeration, paragraph_enumeration

    def _get_paragraph_by_relative_index(self, section_index, paragraph_index):
        found_paragraph = [paragraph for count, paragraph
                           in enumerate(self.master_document.sections[section_index].iter_inner_content())
                           if count == paragraph_index]
        if len(found_paragraph) > 0:
            return found_paragraph[0]
        else:
            raise

    def _get_global_paragraph_index(self, paragraph_content, **kwargs):
        min_paragraph_index = kwargs.get('min_paragraph_index', 0)
        for global_index, paragraph in enumerate(self.master_document.paragraphs):
            if paragraph.text == paragraph_content and global_index >= min_paragraph_index:
                self._logger.debug(f'found \'{paragraph_content}\' at index {global_index}')
                return global_index
        raise AttributeError('paragraph_content not found in document')

    @property
    def file_list(self):
        return self._file_list

    @file_list.setter
    def file_list(self, value):
        if self.master_filename in value:
            value.remove(self.master_filename)
        self._file_list = value

    def _get_config_auto_overwrite(self):
        if self._config:
            self.auto_overwrite_combined_file = self._config.getboolean('UPDATE_LETTER',
                                                                        'auto_overwrite_combined_file')
            self._logger.info(f'config file detected,'
                              f' auto_overwrite_combined_file set to {self.auto_overwrite_combined_file}')
            return
        else:
            self._logger.debug("no config file detected, auto_overwrite_combined_file set to False")
            self.auto_overwrite_combined_file = False

    def file_age(self, file_path, **kwargs):
        """
        This function calculates the age of a file based on its creation time or modification time.

        Parameters:
        - file_path (pathlib.Path): The path to the file.
        - use_create_time (bool, optional): If True, the file creation time will be used to calculate the age. Defaults to False.
        - use_modify_time (bool, optional): If True, the file modification time will be used to calculate the age. Defaults to True.

        Returns:
        - datetime.timedelta: The age of the file as a timedelta object.

        Raises:
        - AttributeError: If both use_create_time and use_modify_time are True or if both are False.

        Example Usage:
        ```
        import datetime
        from pathlib import Path

        # Create a Path object for the file
        file_path = Path('/path/to/file.txt')

        # Calculate the age of the file based on modification time
        age = file_age(file_path, use_modify_time=True)

        # Print the age of the file
        print(age)
        ```
        """
        use_create_time = kwargs.get('use_create_time', False)
        use_modify_time = kwargs.get('use_modify_time', True)
        current_time = datetime.datetime.now().timestamp()
        try:
            if use_create_time and use_modify_time:
                raise AttributeError("use_create_time and use_modify_time cannot both be True")
            if not use_create_time and not use_modify_time:
                raise AttributeError("use_create_time and use_modify_time cannot both be False")
        except AttributeError as e:
            self._logger.error(e, exc_info=True)
            raise e

        if file_path.is_file():
            file_create_time = file_path.stat().st_ctime
            file_modify_time = file_path.stat().st_mtime
        else:
            file_create_time = 0
            file_modify_time = 0
        for x in [{'name': 'file_create_time', 'value': file_create_time},
                  {'name': 'file_modify_time', 'value': file_modify_time}]:
            self._logger.debug(f"time since {x['name']}: {datetime.timedelta(seconds=current_time - x['value'])}")
        if use_modify_time:
            self._logger.debug(f"returning file_modify_time")
            return datetime.timedelta(seconds=current_time - file_modify_time)
        elif use_create_time:
            self._logger.debug(f"returning file_create_time")
            return datetime.timedelta(seconds=current_time - file_create_time)

    def _eval_save_path(self):
        is_not_new_file = self.file_age(self._save_path) > datetime.timedelta(minutes=2)
        overwrite_text = f'The file {self._save_path} already exists, do you want to overwrite it?'
        exists_err_text = f'the file {self._save_path.resolve()} already exists.'
        if is_not_new_file:
            try:
                # TODO: yes to all?
                if self.auto_overwrite_combined_file or questionary.confirm(overwrite_text).ask():
                    self._save_path.unlink()
                else:
                    raise FileExistsError(exists_err_text)
            except FileExistsError as e:
                self._logger.error(e, exc_info=True)
                raise e

    @property
    def save_path(self):
        """
            The `save_path` property is used to access the path where the file is to be saved.

            Returns:
                pathlib.Path: The path of the file to be saved.

            Raises:
                FileExistsError: If the save path already exists and the user chooses not to overwrite it.
            """
        if self._save_path.is_file():
            self._eval_save_path()
        return self._save_path

    @save_path.setter
    def save_path(self, value):
        self._save_path = value

    @staticmethod
    def normalize_child_doc(doc_to_append):
        if isinstance(doc_to_append, (Path, str)):
            doc_to_append = Document(doc_to_append)
        elif isinstance(doc_to_append, Document):
            pass
        return doc_to_append

    def _pre_post_page_breaks(self, doc_to_append, **kwargs):
        page_break_before = kwargs.get('page_break_before', False)
        page_break_after = kwargs.get('page_break_after', False)

        if page_break_after:
            doc_to_append.add_page_break()
        if page_break_before:
            self.composer.doc.add_page_break()

        return doc_to_append

    def _append_to_composer(self, doc_to_append, **kwargs):
        doc_to_append = self._pre_post_page_breaks(
            self.normalize_child_doc(
                doc_to_append), **kwargs
        )
        try:
            self.composer.append(doc_to_append)
        except PackageNotFoundError as e:
            print(e)

    # TODO: turn insert_index into a property and use this?
    def _validate_insert_index(self, **kwargs):
        calculated_index = None
        insert_index_offset = kwargs.get('insert_index_offset', 0)
        placement_paragraph_content = kwargs.get('placement_paragraph_content')
        insert_index = kwargs.get('insert_index')

        if placement_paragraph_content and not insert_index:
            calculated_index = (self._get_global_paragraph_index(
                placement_paragraph_content,
                min_paragraph_index=200) + insert_index_offset)
            insert_index = calculated_index
            self._logger.debug(f"calculated insert_index: {insert_index}")
        elif not placement_paragraph_content and not insert_index:
            raise AttributeError('insert_index or placement_paragraph_content must be provided')
        elif placement_paragraph_content and insert_index:
            raise AttributeError('insert_index and placement_paragraph_content cannot both be provided')

        if insert_index and insert_index != calculated_index:
            self._logger.warning("insert_index_offset is ignored if insert_index is provided")

        return insert_index

    def _insert_in_composer(self, doc_to_insert, insert_index, **kwargs):
        doc_to_insert = self._pre_post_page_breaks(
            self.normalize_child_doc(
                doc_to_insert), **kwargs
        )

        try:
            self.composer.insert(index=insert_index, doc=doc_to_insert)
        except PackageNotFoundError as e:
            print(e)

    def combine_all_docx(self):
        """
            Combines multiple .docx files into a single .docx file.

            This method takes a list of .docx file paths and combines them into a single .docx file.
            It uses the python-docx library to manipulate and merge the files.

            Parameters:
            - self: The instance of the class.
                - type: DocumentMerger
            - return: Returns the path of the combined .docx file.
                - type: str

            How to use:
            1. Create an instance of the DocumentMerger class.
            2. Set the master file using the `set_master_file` method.
            3. Add the files to be merged using the `add_file` method.
            4. Call the `combine_all_docx` method to combine the files.
            5. The method will return the path of the combined file.

            Example:
            ```
            merger = DocumentMerger()
            merger.set_master_file('example-master.docx')
            merger.add_file('example1.docx')
            merger.add_file('example2.docx')
            result = merger.combine_all_docx()
            print(result)
            ```
            """
        for doc in self.file_list:
            self._append_to_composer(doc)
        return self.save_composed_file()

    def save_composed_file(self):
        self.composer.save(self.save_path)
        print(f'combined file saved to {self.save_path}')
        self._logger.info(f'combined file saved to {self.save_path}')
        return self.save_path


class CombineSortedDocx(CombineDocx, FileSorter):
    """
    This class represents a utility for combining multiple sorted Docx files into a single master file.

    class CombineSortedDocx(CombineDocx, FileSorter):
        def __init__(self, master_filename, directory_to_combine, **kwargs):
            Initializes a new instance of the CombineSortedDocx class.

            Parameters:
            - master_filename (str): The name of the master file to be created.
            - directory_to_combine (str): The directory containing the Docx files to be combined.
            - **kwargs: Additional keyword arguments to be passed to the parent class constructors.

            Note:
            - The directory_to_combine parameter should be a valid directory path.

            Example:
                combine_sorted_docx = CombineSortedDocx("master.docx", "/path/to/directory")
        """

    def __init__(self, master_filename, directory_to_combine, **kwargs):
        self._file_list = []
        self._logger = kwargs.get('logger', Logger("dummy_logger"))
        self.directory_to_combine = Path(directory_to_combine)
        FileSorter.__init__(self, self.directory_to_combine, **kwargs)
        super().__init__(master_filename, self.file_list, directory_to_combine, **kwargs)

    @property
    def file_list(self):
        if not self._file_list:
            self._file_list = self.sort_by_last_number()
        return self._file_list

    @file_list.setter
    def file_list(self, value):
        self._file_list = super().file_list


class SingleFileCombine(CombineDocx):
    _DIR_TO_COMBINE_DUMMY_VALUE = ''

    def __init__(self, master_filename, child_file, **kwargs):
        self._child_file = child_file
        super().__init__(master_filename, list(self.child_file),
                         SingleFileCombine._DIR_TO_COMBINE_DUMMY_VALUE, **kwargs)
        self.save_path = Path('./', f'{self.master_filename[:-5]}_combined_file.docx')

    @property
    def child_file(self):
        if isinstance(self._child_file, list):
            self._child_file = self._child_file[0]
        return self._child_file

    def append_and_save(self, **kwargs):
        page_break_before = kwargs.get('page_break_before', False)
        page_break_after = kwargs.get('page_break_after', False)
        self._append_to_composer(self.child_file,
                                 page_break_before=page_break_before,
                                 page_break_after=page_break_after)
        self.save_composed_file()

    def insert_and_save(self, **kwargs):
        page_break_before = kwargs.get('page_break_before', False)
        page_break_after = kwargs.get('page_break_after', False)
        insert_index = self._validate_insert_index(**kwargs)

        self._insert_in_composer(self.child_file, insert_index=insert_index,
                                 page_break_before=page_break_before,
                                 page_break_after=page_break_after)
        self.save_composed_file()


if __name__ == '__main__':
    ...


